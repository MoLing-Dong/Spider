import glob
import os
import re
import subprocess
import sys
import winreg
import zipfile
from random import randint
from time import sleep

import requests
import win32api
from docx import Document
from docx.shared import RGBColor
from fpdf import FPDF
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from tqdm import tqdm

Chapter_names = []


def get_edge_installation_path():  # 获取Edge浏览器路径
    try:
        edge_key = winreg.OpenKey(
            winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\msedge.exe")
        return winreg.QueryValue(edge_key, None)
    except FileNotFoundError:
        return None


def convert_images_to_pdf(image_paths, output_path):
    pdf = FPDF(orientation='L')
    # 添加每张图片到PDF中
    for image_path in image_paths:
        pdf.add_page()
        pdf.image(image_path, 0, 0, pdf.w, pdf.h)  # 根据需要调整图片大小和位置

    # 保存PDF文件
    pdf.output(output_path)


def Get_EdgeVersion():  # 获取Edge浏览器版本号
    try:
        edge_key = r'SOFTWARE\Microsoft\Edge\BLBeacon'
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, edge_key) as key:
            version = winreg.QueryValueEx(key, "version")[0]
            return version
    except Exception as e:
        print(f"Error: {e}")
        return None


def Get_EdgeDriverVersion(file_path):  # 获取Edge浏览器驱动版本号
    try:
        file_info = win32api.GetFileVersionInfo(file_path, "\\")
        ms = file_info['FileVersionMS']
        ls = file_info['FileVersionLS']
        version = f"{win32api.HIWORD(ms)}.{win32api.LOWORD(ms)}.{win32api.HIWORD(ls)}.{win32api.LOWORD(ls)}"
        return version
    except Exception as e:
        print(f"获取文件版本号时出错：{str(e)}")


def Login(Username, Password):  # 学习通登录功能
    driver.get(
        "http://passport2.chaoxing.com/login?fid=&newversion=true&refer=http://i.chaoxing.com")
    Username_element = driver.find_element(By.ID, 'phone')
    Username_element.send_keys(Username)
    sleep(0.3)
    Password_element = driver.find_element(By.ID, 'pwd')
    Password_element.send_keys(Password)
    sleep(0.3)
    loginBtn_element = driver.find_element(By.ID, 'loginBtn')
    loginBtn_element.click()
    print("正在登录中，请稍后...")
    sleep(1)
    for i in range(20):
        if driver.title == "个人空间":
            driver.switch_to.default_content()
            if driver.find_element(By.ID, 'siteName').text == "学校名称":
                return True
        elif i == 20:
            return False
        else:
            sleep(1)


def ClassList():  # 获取课程数量及课程信息
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located((By.ID, 'frame_content')))
    sleep(0.5)
    print("正在获取课程列表中！")
    Home_iframe = driver.find_element(By.ID, 'frame_content')
    driver.switch_to.frame(Home_iframe)
    Classnumber_element = driver.find_elements(
        By.XPATH, "//div[@class='course-info']")  # 获取课程数量
    Tename_element = driver.find_elements(
        By.XPATH, "//div[@class='course-info']/p[2]")  # 获取课程导师数量
    Classname_element = driver.find_elements(
        By.XPATH, "//div[@class='course-info']/p[3]")  # 获取班级名称
    CourseList_element = driver.find_elements(
        By.XPATH, "//span[@class='course-name overHidden2']")  # 获取课程名称数量
    if len(Classnumber_element) > 0:
        print("获取到" + str(len(CourseList_element)) + "门课程：")
        for i in range(len(Classnumber_element)):
            Class_info = str(i) + ". " + CourseList_element[i].text + " " + \
                         Classname_element[i].text + " 科任老师：" + Tename_element[i].text
            print(Class_info)
        Classnumber = input("请输入你要学习的课程序号：")
        if Classnumber != "":
            Classname = CourseList_element[int(Classnumber)].text
            CourseList_element[int(Classnumber)].click()
            print("正在进入：" + Classname + "课程中，请稍后....")
            sleep(1)
            Wins = driver.window_handles
            if len(Wins) > 1:
                driver.switch_to.window(Wins[1])
                sleep(1)
                for i in range(20):
                    if driver.title == Classname:
                        return True
                    elif i == 10:
                        return False
                    else:
                        sleep(1)
        else:
            print("错误：课程序号不能为空！")
            return False
    else:
        print("获取课程失败！")
        return False


def Chapter():  # 获取章节数量及章节信息
    global EnterSection_number
    Finsh, Unfinsh = 0, 0
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//a[@data-url='/mooc2-ans/mycourse/studentcourse']")))
    sleep(0.5)
    Chapter_element = driver.find_element(
        By.XPATH, "//a[@data-url='/mooc2-ans/mycourse/studentcourse']")
    if Chapter_element.text == "章节":
        print("正在获取章节信息中，请稍等....")
        Chapter_element.click()
        wait = WebDriverWait(driver, 10)  # 等待网页加载完成
        wait.until(EC.presence_of_element_located((By.ID, 'frame_content-zj')))
        sleep(0.5)
        Chapter_iframe = driver.find_element(By.ID, 'frame_content-zj')
        driver.switch_to.frame(Chapter_iframe)
        Chapternumber_element = driver.find_elements(
            By.XPATH, "//div[@class='catalog_title chapter_Thats_bnt']/div[3]/span")  # 获取章节名称数量
        Sectionnumber_element = driver.find_elements(
            By.XPATH, "//div[@class='catalog_task']/div")  # 获取小节名称数量
        Sectiontitle_element = driver.find_elements(
            By.XPATH, "//div[@class='catalog_title']/div[2]")  # 获取小节名称
        if len(Chapternumber_element) > 0:
            for i in range(len(Sectionnumber_element)):
                Statusnumber = Sectionnumber_element[i].find_elements(
                    By.XPATH, "span")
                Chapter_names.append(Sectiontitle_element[i].text)
                if len(Statusnumber) == 1:
                    Finsh += 1
                    print(str(i) + "  " +
                          Sectiontitle_element[i].text + "  完成状态：完成")
                elif len(Statusnumber) == 2:
                    Unfinsh += 1
                    print(str(i) + "  " +
                          Sectiontitle_element[i].text + "  完成状态：未完成")
            print("已获取到：" + str(len(Chapternumber_element)) +
                  "个章节 " + str(len(Sectionnumber_element)) + "小节,任务点已完成：" + str(
                Finsh) + "个小节，未完成：" + str(Unfinsh) + "个小节")
            EnterSection_number = input("请输入要进入的小节：")
            if EnterSection_number != "":
                Sectiontitle_element[int(EnterSection_number)].click()
                print("正在进入小节中，请稍等...")
                sleep(1)
                for i in range(20):
                    if driver.title == "学生学习页面":
                        return True
                    else:
                        sleep(1)
            else:
                print("错误：小节序号不能为空！")
                return False

        else:
            print("获取章节数失败！")
            return False
    else:
        print("获取章节定位失败！")
        return False


def Examination():  # 爬取考试题目试卷及正确答案
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//div[@class='nav-content   stuNavigationList']")))
    print("正在进入考试中！")
    考试标签 = driver.find_elements(By.XPATH, "//li[@dataname='ks']")
    if len(考试标签) > 0:
        driver.find_element(By.XPATH, "//li[@dataname='ks']").click()
        wait = WebDriverWait(driver, 10)  # 等待网页加载完成
        wait.until(EC.presence_of_element_located(
            (By.XPATH, "//iframe[@id='frame_content-ks']")))
        iframe = driver.find_element(
            By.XPATH, "//iframe[@id='frame_content-ks']")
        driver.switch_to.frame(iframe)
        考试数量 = driver.find_elements(
            By.XPATH, "//div[@class='bottomList']/ul/li")
        if len(考试数量) > 0:
            for i in range(len(考试数量)):
                driver.find_element(
                    By.XPATH, f"//div[@class='bottomList']/ul/li[{i + 1}]/div").click()
                Winhandle = driver.window_handles
                if len(Winhandle) == 3:
                    driver.switch_to.window(Winhandle[2])
                    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
                    wait.until(EC.presence_of_element_located(
                        (By.XPATH, "//div[@class='result_Main']/p")))
                    是否可以查看试卷 = driver.find_elements(
                        By.XPATH, f"//div[@class='result_Main']/p/a")
                    if len(是否可以查看试卷) > 0:
                        print("开始爬取试卷，请稍等！")
                        doc = Document()
                        driver.find_element(
                            By.XPATH, f"//div[@class='result_Main']/p/a").click()
                        wait = WebDriverWait(driver, 10)  # 等待网页加载完成
                        wait.until(EC.presence_of_element_located(
                            (By.XPATH, "//head/title")))
                        试卷标题 = driver.find_element(
                            By.XPATH, "//h2[@class='mark_title']").text
                        试卷题目数量标题 = driver.find_element(
                            By.XPATH, "//div[@class='infoHead fl']/span[1]").text
                        试卷满分标题 = driver.find_element(
                            By.XPATH, "//div[@class='infoHead fl']/span[2]").text
                        doc.add_paragraph(试卷标题)
                        doc.add_paragraph(试卷题目数量标题)
                        doc.add_paragraph(试卷满分标题)
                        div_element = driver.find_element(
                            By.XPATH, "//div[@class='mark_item ans-cc']")
                        Child_elements = div_element.find_elements(
                            By.XPATH, "./*")
                        h2, div = 1, 1
                        for Child_element in Child_elements:
                            if Child_element.tag_name == "h2":  # 题目类型
                                题目类型标题 = driver.find_element(
                                    By.XPATH, f"//div[@class='mark_item ans-cc']/h2[{h2}]").text
                                doc.add_paragraph(题目类型标题)
                                h2 += 1
                            elif Child_element.tag_name == "div":  # 题目
                                题目 = driver.find_element(
                                    By.XPATH, f"//div[@class='mark_item ans-cc']/div[{div}]/h3").text
                                题目分数及类型 = driver.find_element(
                                    By.XPATH, f"//div[@class='mark_item ans-cc']/div[{div}]/h3/span").text
                                题目分数 = re.findall(r'\d+\.\d+分', 题目分数及类型)[0]
                                doc.add_paragraph(题目)
                                是否有选项 = driver.find_elements(
                                    By.XPATH, f"//div[@class='mark_item ans-cc']/div[{div}]/ul")
                                if len(是否有选项) > 0:
                                    选项数量 = driver.find_elements(
                                        By.XPATH, f"//div[@class='mark_item ans-cc']/div[{div}]/ul/li")
                                    for 选项计数 in range(len(选项数量)):
                                        选项 = driver.find_element(
                                            By.XPATH,
                                            f"//div[@class='mark_item ans-cc']/div[{div}]/ul/li[{选项计数 + 1}]").text
                                        doc.add_paragraph(选项)
                                答题标签数量 = driver.find_elements(
                                    By.XPATH, f"//div[@class='mark_item ans-cc']/div[{div}]/div/div")
                                if len(答题标签数量) > 0:
                                    for 标签计数 in range(len(答题标签数量)):
                                        标签名称 = driver.find_element(
                                            By.XPATH,
                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[{标签计数 + 1}]").get_attribute(
                                            "Class")
                                        if 标签名称 == "mark_score":
                                            选择是否存在对错 = driver.find_elements(
                                                By.XPATH,
                                                f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[{标签计数 + 1}]/div")
                                            if len(选择是否存在对错) > 0:
                                                for 选择对错计次 in range(len(选择是否存在对错)):
                                                    选择对错标签 = driver.find_element(
                                                        By.XPATH,
                                                        f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[{标签计数 + 1}]/div[{选择对错计次 + 1}]").get_attribute(
                                                        "Class")
                                                    if 选择对错标签 == "totalScore fr":
                                                        答题得分 = driver.find_element(
                                                            By.XPATH,
                                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[{标签计数 + 1}]/div[{选择对错计次 + 1}]").text
                                                        答案 = driver.find_element(
                                                            By.XPATH,
                                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[1]/span").text
                                                        if 答题得分 == 题目分数:
                                                            green_run = doc.add_paragraph().add_run(答案)
                                                            green_run.font.color.rgb = RGBColor(
                                                                0, 255, 0)
                                                        else:
                                                            答案数量 = driver.find_elements(
                                                                By.XPATH,
                                                                f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[1]/span")
                                                            if len(答案数量) > 1:  # 存在正确答案
                                                                for 答案数量计次 in range(len(答案数量)):
                                                                    答案标签 = driver.find_element(
                                                                        By.XPATH,
                                                                        f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[1]/span[{答案数量计次 + 1}]").get_attribute(
                                                                        "Class")
                                                                    if 答案标签 == "colorGreen marginRight40 fl":
                                                                        答案 = driver.find_element(
                                                                            By.XPATH,
                                                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div[1]/span").text
                                                                        green_run = doc.add_paragraph().add_run(答案)
                                                                        green_run.font.color.rgb = RGBColor(
                                                                            0, 255, 0)
                                                            else:
                                                                red_run = doc.add_paragraph().add_run("我的答案：")
                                                                red_run.font.color.rgb = RGBColor(
                                                                    255, 0, 0)
                                        elif 标签名称 == None:  # 填空题
                                            填空是否存在对错 = driver.find_elements(
                                                By.XPATH, f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl")
                                            if len(填空是否存在对错) > 1:  # 给出正确答案
                                                for 填空对错计次 in range(len(填空是否存在对错)):
                                                    填空对错标签 = driver.find_element(
                                                        By.XPATH,
                                                        f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl[{填空对错计次 + 1}]").get_attribute(
                                                        "Class")
                                                    if 填空对错标签 == "mark_fill colorGreen":
                                                        答案标题 = driver.find_element(
                                                            By.XPATH,
                                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl[{填空对错计次 + 1}]/dt").text
                                                        答案数量 = driver.find_elements(
                                                            By.XPATH,
                                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl[{填空对错计次 + 1}]/dd")
                                                        green_run = doc.add_paragraph().add_run(答案标题)
                                                        green_run.font.color.rgb = RGBColor(
                                                            0, 255, 0)
                                                        for 答案计次 in range(len(答案数量)):
                                                            答案 = driver.find_element(
                                                                By.XPATH,
                                                                f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl[{填空对错计次 + 1}]/dd[{答案计次 + 1}]").text
                                                            green_run = doc.add_paragraph().add_run(答案)
                                                            green_run.font.color.rgb = RGBColor(
                                                                0, 255, 0)
                                            elif 0 < len(填空是否存在对错) <= 1:  # 仅有自己的答题答案
                                                答案标题 = driver.find_element(
                                                    By.XPATH,
                                                    f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl/dt/i").text
                                                答题标签数量 = driver.find_elements(
                                                    By.XPATH,
                                                    f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl/dt/div/div")
                                                for 答题标签计次 in range(len(答题标签数量)):
                                                    答题标签 = driver.find_element(
                                                        By.XPATH,
                                                        f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl/dt/div/div[{答题标签计次 + 1}]").get_attribute(
                                                        "Class")
                                                    if 答题标签 == "totalScore fr":
                                                        答题得分 = driver.find_element(
                                                            By.XPATH,
                                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl/dt/div/div[{答题标签计次 + 1}]").text
                                                        答案数量 = driver.find_elements(
                                                            By.XPATH,
                                                            f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl/dd")
                                                        if 答题得分 == 题目分数:
                                                            green_run = doc.add_paragraph().add_run(答案标题)
                                                            green_run.font.color.rgb = RGBColor(
                                                                0, 255, 0)
                                                            for 答案计次 in range(len(答案数量)):
                                                                答案 = driver.find_element(
                                                                    By.XPATH,
                                                                    f"//div[@class='mark_item ans-cc']/div[{div}]/div/div/dl/dd[{答案计次 + 1}]/span").text
                                                                green_run = doc.add_paragraph().add_run(答案)
                                                                green_run.font.color.rgb = RGBColor(
                                                                    0, 255, 0)
                                                        else:
                                                            red_run = doc.add_paragraph().add_run(答案标题)
                                                            red_run.font.color.rgb = RGBColor(
                                                                255, 0, 0)
                                                            for 答案计次 in range(len(答案数量)):
                                                                red_run = doc.add_paragraph().add_run(
                                                                    f"({答案计次 + 1})")
                                                                red_run.font.color.rgb = RGBColor(
                                                                    255, 0, 0)
                                div += 1
                        Save_Path = input("请输入保存到的位置：(0.运行目录 1.系统桌面 2.自定义路径)：")
                        if Save_Path == "0":
                            root_path = os.getcwd()
                            doc.save(root_path + "\\" + 试卷标题 + ".docx")
                            print("保存路径：" + root_path + "\\" + 试卷标题 + ".docx")
                        elif Save_Path == "1":
                            root_path = os.path.join(
                                os.path.expanduser("~"), "Desktop")
                            doc.save(root_path + "\\" + 试卷标题 + ".docx")
                            print("保存路径：" + root_path + "\\" + 试卷标题 + ".docx")
                        elif Save_Path == "2":
                            root_path = input("请输入保存的路径：")
                            root_path = root_path.rstrip("/")
                            doc.save(root_path + "\\" + 试卷标题 + ".docx")
                            print("保存路径：" + root_path + "\\" + 试卷标题 + ".docx")
                        print("爬取考试试卷成功！")
                        sleep(3)
                        driver.close()
                        driver.switch_to.window(Winhandle[1])
                        driver.switch_to.frame(iframe)
                    else:
                        print("本考试卷考后不允许查看！")
                        driver.close()
                        driver.switch_to.window(Winhandle[1])
                        driver.switch_to.frame(iframe)


def Function():  # 功能选择汇总
    print("脚本功能：")
    print("    1.完成当前课程任务点")
    print("    2.刷取当前课程视频时长")
    print("    3.刷取当前课程阅读时长")
    print("    4.爬取指定小节及以下PPT课件")
    print("    5.爬取指定小节PPT课件")
    print("    6.爬取指定小节及以下视频")
    print("    7.爬取指定小节视频")
    print("    8.爬取当前课程考试题目")
    print("    9.爬取当前课程练习题题目")
    Function_number = input("请输入实行的模式：")
    if Function_number == "1":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Finsh_task()
        else:
            print("进入章节失败！")
            input("")
    elif Function_number == "2":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Player_video()
        else:
            print("进入章节失败！")
            input("")
    elif Function_number == "3":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Read_ppt()
        else:
            print("进入章节失败！")
            input("")
    elif Function_number == "4":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Downlaod_PPts()
        else:
            print("进入章节失败！")
            input("")
    elif Function_number == "5":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Download_PPt()
        else:
            print("进入章节失败！")
            input("")
    elif Function_number == "6":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Download_videos()
        else:
            print("进入章节失败！")
            input("")
    elif Function_number == "7":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Download_video()
        else:
            print("进入章节失败！")
            input("")
    elif Function_number == "8":
        Examination()
    elif Function_number == "9":
        Enterchapter = Chapter()
        if Enterchapter == True:
            print("进入章节成功！")
            Downlaod_Exercises()
        else:
            print("进入章节失败！")
            input("")
    else:
        print("现仅支持八种模式！")
        return


def Downlaod_Exercises():
    pass


def Finsh_task():
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//iframe[@id='iframe']")))
    Chapter_number = driver.find_elements(
        By.XPATH, "//div[@class='onetoone posCatalog']/ul/li")
    if len(Chapter_number) > 0:
        print("当前课程具有" + str(len(Chapter_number)) + "章节")
        for i in range(len(Chapter_number)):
            Litterchapter_number = driver.find_elements(
                By.XPATH, f"//div[@class='onetoone posCatalog']/ul/li[{i + 1}]/div[2]/ul/li")
            if len(Litterchapter_number) > 0:
                for a in range(len(Litterchapter_number)):
                    Title = driver.find_elements(
                        By.XPATH,
                        f"//div[@class='onetoone posCatalog']/ul/li[{i + 1}]/div[2]/ul/li[{a + 1}]/div/span[1]")
                    for b in range(len(Title)):
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Title[b])
                        wait = WebDriverWait(driver, 10)  # 等待网页加载完成
                        wait.until(EC.presence_of_element_located(
                            (By.XPATH, "//iframe[@id='iframe']")))
                        Task_title = driver.find_element(
                            By.XPATH, "//div[@class='prev_title']").get_attribute("title")
                        iframe = driver.find_element(
                            By.XPATH, "//iframe[@id='iframe']")
                        driver.switch_to.frame(iframe)
                        Task_finsh = driver.find_elements(
                            By.XPATH, "//div[@class='ans-attach-ct ans-job-finished']")
                        Task_unfinsh = driver.find_elements(
                            By.XPATH, "//div[@class='ans-attach-ct']/iframe")
                        Task_number = len(Task_finsh) + len(Task_unfinsh)
                        print(Task_title + " 具有：" + str(Task_number) + "个任务点,已完成：" +
                              str(len(Task_finsh)) + " 未完成：" + str(len(Task_unfinsh)))
                        if len(Task_unfinsh) > 0:
                            for c in range(len(Task_unfinsh)):
                                driver.execute_script(
                                    "arguments[0].scrollIntoView();", Task_unfinsh[c])
                                Task_data = Task_unfinsh[c].get_attribute(
                                    "data")
                                type_pattern = r'"type":"(.*?)"'
                                Task_type = re.findall(type_pattern, Task_data)
                                if Task_type[0] == ".pdf":
                                    driver.switch_to.frame(Task_unfinsh[c])
                                    sleep(0.5)
                                    Pdf_iframe = driver.find_element(
                                        By.XPATH, "//iframe[@id='panView']")
                                    driver.switch_to.frame(Pdf_iframe)
                                    sleep(0.5)
                                    Pdf_number = driver.find_elements(
                                        By.XPATH, "//div[@class='fileBox']/ul/li")
                                    if len(Pdf_number) > 0:
                                        print("开始阅读PPT：")
                                        for d in range(len(Pdf_number)):
                                            Pdf_now = driver.find_element(
                                                By.XPATH, f"//div[@class='fileBox']/ul/li[{d + 1}]")
                                            driver.execute_script(
                                                "arguments[0].scrollIntoView();", Pdf_now)
                                            if b + 1 == len(Pdf_number):
                                                print("\r" + str(d + 1) +
                                                      " / " + str(len(Pdf_number)))
                                                driver.switch_to.default_content()
                                                sleep(1)
                                                driver.switch_to.frame(iframe)
                                            else:
                                                print("\r" + str(d + 1) +
                                                      " / " + str(len(Pdf_number)), end="")
                                            sleep(randint(30, 60))
                                        print("阅读PPT完毕！")
                                elif Task_type[0] == ".mp4":
                                    driver.switch_to.frame(Task_unfinsh[c])
                                    sleep(0.5)
                                    Tip_element = driver.find_element(
                                        By.XPATH, "//div[@id='reader']")
                                    sleep(0.5)
                                    driver.execute_script(
                                        "arguments[0].scrollIntoView();", Tip_element)
                                    sleep(1)
                                    Videoplay_element = driver.find_element(
                                        By.XPATH, "//button[@class='vjs-big-play-button']")
                                    sleep(0.5)
                                    Videoplay_element.click()
                                    sleep(1)
                                    while True:
                                        Videoend_time = driver.find_element(
                                            By.XPATH, "//span[@class='vjs-duration-display']").get_attribute(
                                            "textContent")
                                        if Videoend_time != "0:00":
                                            break
                                        else:
                                            sleep(0.5)
                                    print("正在播放视频，播放进度：")
                                    while True:
                                        Videonow_time = driver.find_element(
                                            By.XPATH, "//span[@class='vjs-current-time-display']").get_attribute(
                                            "textContent")
                                        if Videonow_time == Videoend_time:
                                            print("\r" + Videonow_time +
                                                  " / " + Videoend_time)
                                            break
                                        else:
                                            print("\r" + Videonow_time +
                                                  " / " + Videoend_time, end="")
                                            sleep(1)
                            print("已完成本章节未完成的任务点！")
                            driver.switch_to.default_content()
                            sleep(1)
                            Next_button = driver.find_elements(
                                By.XPATH, "//div[@id='prevNextFocusNext']")
                            if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                                driver.execute_script(
                                    "arguments[0].scrollIntoView();", Next_button[0])
                                sleep(1)
                                if len(Next_button) > 0:
                                    Next_button[0].click()
                                    print("下一节")
                                sleep(10)
                            else:
                                print("当前课程任务点已全部完成！")
                        else:
                            print("未发现未完成的任务点！")
                            driver.switch_to.default_content()
                            sleep(1)
                            Next_button = driver.find_elements(
                                By.XPATH, "//div[@id='prevNextFocusNext']")
                            if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                                driver.execute_script(
                                    "arguments[0].scrollIntoView();", Next_button[0])
                                sleep(1)
                                if len(Next_button) > 0:
                                    Next_button[0].click()
                                    print("下一节")
                                sleep(10)
                            else:
                                print("当前课程任务点已全部完成！")
            else:
                print("未发现小节！")


def Player_video():  # 刷视频时长功能
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//iframe[@id='iframe']")))
    End_Section = False
    while True:
        Video_Fistiframe = driver.find_element(
            By.XPATH, "//iframe[@id='iframe']")
        driver.switch_to.frame(Video_Fistiframe)  # 切换至第一层
        sleep(1)
        Video_number = driver.find_elements(
            By.XPATH, "//iframe[@class='ans-attach-online ans-insertvideo-online']")
        if len(Video_number) > 0:
            print("发现" + str(len(Video_number)) + "个视频任务")
            for i in range(len(Video_number)):
                driver.switch_to.frame(Video_number[i])
                sleep(0.5)
                Tip_element = driver.find_element(
                    By.XPATH, "//div[@id='reader']")
                sleep(0.5)
                driver.execute_script(
                    "arguments[0].scrollIntoView();", Tip_element)
                sleep(1)
                Videoplay_element = driver.find_element(
                    By.XPATH, "//button[@class='vjs-big-play-button']")
                sleep(0.5)
                Videoplay_element.click()
                sleep(1)
                while True:
                    Videoend_time = driver.find_element(
                        By.XPATH, "//span[@class='vjs-duration-display']").get_attribute("textContent")
                    if Videoend_time != "0:00":
                        break
                    else:
                        sleep(0.5)
                print("正在播放第" + str(i + 1) + "视频，播放进度：")
                while True:
                    Videonow_time = driver.find_element(
                        By.XPATH, "//span[@class='vjs-current-time-display']").get_attribute("textContent")
                    if Videonow_time == Videoend_time:
                        print("\r" + Videonow_time +
                              " / " + Videoend_time)
                        break
                    else:
                        print("\r" + Videonow_time +
                              " / " + Videoend_time, end="")
                        sleep(1)
                driver.switch_to.parent_frame()

            if End_Section == False:
                sleep(0.5)
                driver.switch_to.default_content()
                sleep(1)
                Next_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusNext']")
                if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Next_button[0])
                    sleep(1)
                    if len(Next_button) > 0:
                        Next_button[0].click()
                        print("下一节")
                    sleep(10)
                else:
                    Prev_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusPrev']")
                    if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Prev_button[0])
                        sleep(1)
                        if len(Prev_button) > 0:
                            Prev_button[0].click()
                            print("上一节")
                        sleep(10)
                        End_Section = True
            elif End_Section == True:
                sleep(0.5)
                driver.switch_to.default_content()
                sleep(1)
                Prev_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusPrev']")
                if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Prev_button[0])
                    sleep(1)
                    if len(Prev_button) > 0:
                        Prev_button[0].click()
                        print("上一节")
                    sleep(10)
                else:
                    Next_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusNext']")
                    if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Next_button[0])
                        sleep(1)
                        if len(Next_button) > 0:
                            Next_button[0].click()
                            print("下一节")
                        sleep(10)
                        End_Section = False
        else:
            print("未发现视频任务！")
            if End_Section == False:
                sleep(0.5)
                driver.switch_to.default_content()
                sleep(1)
                Next_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusNext']")
                if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Next_button[0])
                    sleep(1)
                    if len(Next_button) > 0:
                        Next_button[0].click()
                        print("下一节")
                    sleep(10)
                else:
                    Prev_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusPrev']")
                    if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Prev_button[0])
                        sleep(1)
                        if len(Prev_button) > 0:
                            Prev_button[0].click()
                            print("上一节")
                        sleep(10)
                    End_Section = True
            elif End_Section == True:
                sleep(0.5)
                driver.switch_to.default_content()
                sleep(1)
                Prev_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusPrev']")
                if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Prev_button[0])
                    sleep(1)
                    if len(Prev_button) > 0:
                        Prev_button[0].click()
                        print("上一节")
                    sleep(10)
                else:
                    sleep(0.5)
                    driver.switch_to.default_content()
                    sleep(1)
                    Next_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusNext']")
                    if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Next_button[0])
                        sleep(1)
                        if len(Next_button) > 0:
                            Next_button[0].click()
                            print("下一节")
                        sleep(10)
                    End_Section = False


def Read_ppt():  # 刷PPT阅读时长功能
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//iframe[@id='iframe']")))
    End_Section = False
    while True:
        Firstiframe_Readppt = driver.find_element(
            By.XPATH, "//iframe[@id='iframe']")
        driver.switch_to.frame(Firstiframe_Readppt)
        sleep(1)
        Secondiframe_Readppt = driver.find_elements(
            By.XPATH, "//iframe[@class='ans-attach-online insertdoc-online-ppt']")
        if len(Secondiframe_Readppt) > 0:
            for i in range(len(Secondiframe_Readppt)):
                driver.switch_to.frame(Secondiframe_Readppt[i])
                sleep(0.5)
                Readppt_number = driver.find_elements(
                    By.XPATH, "//iframe[@id='panView']")
                if len(Readppt_number) > 0:
                    for i in range(len(Readppt_number)):
                        driver.switch_to.frame(Readppt_number[i])
                        sleep(0.5)
                        Tip_element = driver.find_element(
                            By.XPATH, "//div[@class='fileBox']")
                        sleep(0.5)
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Tip_element)
                        sleep(1)
                        Read_ppt_number = driver.find_elements(
                            By.XPATH, "//div[@class='fileBox']/ul/li")
                        if len(Read_ppt_number) > 0:
                            print("开始阅读PPT:")
                            for i in range(len(Read_ppt_number)):
                                Now_ppt = Read_ppt_number[i].find_element(
                                    By.XPATH, f"//div[@class='fileBox']/ul/li[{i + 1}]")
                                sleep(1)
                                driver.execute_script(
                                    "arguments[0].scrollIntoView();", Now_ppt)
                                if i + 1 == len(Read_ppt_number):
                                    print("\r" + str(i + 1) +
                                          " / " + str(len(Read_ppt_number)))
                                else:
                                    print("\r" + str(i + 1) +
                                          " / " + str(len(Read_ppt_number)), end="")
                                    sleep(randint(30, 60))
                            if End_Section == False:
                                driver.switch_to.default_content()
                                sleep(1)
                                Next_button = driver.find_elements(
                                    By.XPATH, "//div[@id='prevNextFocusNext']")
                                if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                                    driver.execute_script(
                                        "arguments[0].scrollIntoView();", Next_button[0])
                                    sleep(1)
                                    if len(Next_button) > 0:
                                        Next_button[0].click()
                                        print("下一节")
                                    sleep(10)
                                else:
                                    Prev_button = driver.find_elements(
                                        By.XPATH, "//div[@id='prevNextFocusPrev']")
                                    if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                                        driver.execute_script(
                                            "arguments[0].scrollIntoView();", Prev_button[0])
                                        sleep(1)
                                        if len(Prev_button) > 0:
                                            Prev_button[0].click()
                                            print("上一节")
                                        sleep(10)
                                        End_Section = True
                            elif End_Section == True:
                                driver.switch_to.default_content()
                                sleep(1)
                                Prev_button = driver.find_elements(
                                    By.XPATH, "//div[@id='prevNextFocusPrev']")
                                if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                                    driver.execute_script(
                                        "arguments[0].scrollIntoView();", Prev_button[0])
                                    sleep(1)
                                    if len(Prev_button) > 0:
                                        Prev_button[0].click()
                                        print("上一节")
                                    sleep(10)
                                else:
                                    Next_button = driver.find_elements(
                                        By.XPATH, "//div[@id='prevNextFocusNext']")
                                    if len(Next_button) > 0 and Next_button[0].get_attribute(
                                            "style") != "display: none;":
                                        driver.execute_script(
                                            "arguments[0].scrollIntoView();", Next_button[0])
                                        sleep(1)
                                        if len(Next_button) > 0:
                                            Next_button[0].click()
                                            print("下一节")
                                        sleep(10)
                                        End_Section = False
                        else:
                            print("未发现PPT！")
        else:
            print("未发现PPT任务！")
            if End_Section == False:
                sleep(0.5)
                driver.switch_to.default_content()
                sleep(1)
                Next_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusNext']")
                if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Next_button[0])
                    sleep(1)
                    if len(Next_button) > 0:
                        Next_button[0].click()
                        print("下一节")
                    sleep(10)
                else:
                    Prev_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusPrev']")
                    if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Prev_button[0])
                        sleep(1)
                        if len(Prev_button) > 0:
                            Prev_button[0].click()
                            print("上一节")
                        sleep(10)
                        End_Section = True
            elif End_Section == True:
                sleep(0.5)
                driver.switch_to.default_content()
                sleep(1)
                Prev_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusPrev']")
                if len(Prev_button) and Prev_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Prev_button[0])
                    sleep(1)
                    if len(Prev_button) > 0:
                        Prev_button[0].click()
                        print("上一节")
                    sleep(10)
                else:
                    Next_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusNext']")
                    if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Next_button[0])
                        sleep(1)
                        if len(Next_button) > 0:
                            Next_button[0].click()
                            print("下一节")
                        sleep(10)
                        End_Section = False


def Downlaod_PPts():  # 爬取当前小节及以下PPT课件
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//iframe[@id='iframe']")))
    if len(Chapter_names) > 0:
        root_path = os.getcwd()
        if os.path.exists(root_path + r"\学习通PPT课件") == False:
            os.mkdir(root_path + r"\学习通PPT课件")
        if os.path.exists(root_path + r"\学习通PPT课件\Iamge") == False:
            os.mkdir(root_path + r"\学习通PPT课件\Iamge")
        if os.path.exists(root_path + r"\学习通PPT课件\Pdf文件") == False:
            os.mkdir(root_path + r"\学习通PPT课件\Pdf文件")
        image_paths = []
        for a in range(len(Chapter_names)):
            Firstiframe_Readppt = driver.find_element(
                By.XPATH, "//iframe[@id='iframe']")
            driver.switch_to.frame(Firstiframe_Readppt)
            sleep(1)
            Secondiframe_Readppt = driver.find_elements(
                By.XPATH, "//iframe[@class='ans-attach-online insertdoc-online-ppt']")
            if len(Secondiframe_Readppt) > 0:
                for i in range(len(Secondiframe_Readppt)):
                    driver.switch_to.frame(Secondiframe_Readppt[i])
                    sleep(0.5)
                    Readppt_number = driver.find_elements(
                        By.XPATH, "//iframe[@id='panView']")
                    if len(Readppt_number) > 0:
                        for i in range(len(Readppt_number)):
                            driver.switch_to.frame(Readppt_number[i])
                            sleep(0.5)
                            Tip_element = driver.find_element(
                                By.XPATH, "//div[@class='fileBox']")
                            sleep(0.5)
                            driver.execute_script(
                                "arguments[0].scrollIntoView();", Tip_element)
                            sleep(1)
                            Read_ppt_number = driver.find_elements(
                                By.XPATH, "//div[@class='fileBox']/ul/li")
                            if len(Read_ppt_number) > 0:
                                for i in range(len(Read_ppt_number)):
                                    Url_ppt = Read_ppt_number[i].find_element(
                                        By.XPATH, f"//div[@class='fileBox']/ul/li[{i + 1}]/img").get_attribute("src")
                                    print("开始爬取第" + str(i + 1) + "张")
                                    response = requests.get(
                                        url=Url_ppt, stream=True)
                                    with open(root_path + f"\\学习通PPT课件\\Iamge\\image_{i}.png", "wb") as fp:
                                        image_paths.append(
                                            root_path + f"\\学习通PPT课件\\Iamge\\image_{i}.png")
                                        fp.write(response.content)
                            else:
                                print("未发现PPT！")
                    else:
                        print("未发现PPT任务！")
                if len(image_paths) > 0:
                    convert_images_to_pdf(
                        image_paths, root_path + f"\\学习通PPT课件\\Pdf文件\\{Chapter_names[a]}.pdf")
                    print(f"合成PDF文件成功，文件名称：{Chapter_names[a]}.pdf")
                    image_files = glob.glob(os.path.join(
                        root_path + "\\学习通PPT课件\\Iamge", '*.png'))
                    for file in image_files:
                        os.remove(file)
                    image_paths.clear()
                    driver.switch_to.default_content()
                    sleep(1)
                    Next_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusNext']")
                    if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Next_button[0])
                        sleep(1)
                        if len(Next_button) > 0:
                            Next_button[0].click()
                            print("下一节")
                        sleep(10)
                    else:
                        print("爬取PPT课件完成！文件保存路径：" +
                              root_path + "\\学习通PPT课件\\Pdf文件")

                else:
                    print("错误：合成PDF文件失败！")
                    break
            else:
                print("未发现PPT！")
                image_paths.clear()
                driver.switch_to.default_content()
                sleep(1)
                Next_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusNext']")
                if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Next_button[0])
                    sleep(1)
                    if len(Next_button) > 0:
                        Next_button[0].click()
                        print("下一节")
                    sleep(10)
    else:
        print("获取章节失败！")


def Download_PPt():  # 爬取当前小节PPT课件
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//iframe[@id='iframe']")))
    if len(Chapter_names) > 0:
        root_path = os.getcwd()
        if os.path.exists(root_path + r"\学习通PPT课件") == False:
            os.mkdir(root_path + r"\学习通PPT课件")
        if os.path.exists(root_path + r"\学习通PPT课件\Iamge") == False:
            os.mkdir(root_path + r"\学习通PPT课件\Iamge")
        if os.path.exists(root_path + r"\学习通PPT课件\Pdf文件") == False:
            os.mkdir(root_path + r"\学习通PPT课件\Pdf文件")
        image_paths = []
        Firstiframe_Readppt = driver.find_element(
            By.XPATH, "//iframe[@id='iframe']")
        driver.switch_to.frame(Firstiframe_Readppt)
        sleep(1)
        Secondiframe_Readppt = driver.find_elements(
            By.XPATH, "//iframe[@class='ans-attach-online insertdoc-online-ppt']")
        if len(Secondiframe_Readppt) > 0:
            for i in range(len(Secondiframe_Readppt)):
                driver.switch_to.frame(Secondiframe_Readppt[i])
                sleep(0.5)
                Readppt_number = driver.find_elements(
                    By.XPATH, "//iframe[@id='panView']")
                if len(Readppt_number) > 0:
                    for i in range(len(Readppt_number)):
                        driver.switch_to.frame(Readppt_number[i])
                        sleep(0.5)
                        Tip_element = driver.find_element(
                            By.XPATH, "//div[@class='fileBox']")
                        sleep(0.5)
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Tip_element)
                        sleep(1)
                        Read_ppt_number = driver.find_elements(
                            By.XPATH, "//div[@class='fileBox']/ul/li")
                        if len(Read_ppt_number) > 0:
                            for i in range(len(Read_ppt_number)):
                                Url_ppt = Read_ppt_number[i].find_element(
                                    By.XPATH, f"//div[@class='fileBox']/ul/li[{i + 1}]/img").get_attribute("src")
                                print("开始爬取第" + str(i + 1) + "张")
                                response = requests.get(
                                    url=Url_ppt, stream=True)
                                with open(root_path + f"\\学习通PPT课件\\Iamge\\image_{i}.png", "wb") as fp:
                                    image_paths.append(
                                        root_path + f"\\学习通PPT课件\\Iamge\\image_{i}.png")
                                    fp.write(response.content)
                        else:
                            print("未发现PPT！")
                else:
                    print("未发现PPT任务！")
        if len(image_paths) > 0:
            convert_images_to_pdf(
                image_paths, root_path + f"\\学习通PPT课件\\Pdf文件\\{Chapter_names[int(EnterSection_number)]}.pdf")
            print(
                f"合成PDF文件成功，文件名称：{Chapter_names[int(EnterSection_number)]}.pdf")
            image_files = glob.glob(os.path.join(
                root_path + "\\学习通PPT课件\\Iamge", '*.png'))
            for file in image_files:
                os.remove(file)
            image_paths.clear()
            print("爬取PPT课件完成！文件保存路径：" +
                  root_path + "\\学习通PPT课件\\Pdf文件")
        else:
            print("错误：合成PDF文件失败！")
    else:
        print("获取章节失败！")


def Download_videos():  # 爬取当前小节及以下视频
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//iframe[@id='iframe']")))
    if len(Chapter_names) > 0:
        Video_name = []
        Video_url = []
        Video_header = []
        root_path = os.getcwd()
        if os.path.exists(root_path + r"\学习通章节视频") == False:
            os.mkdir(root_path + r"\学习通章节视频")
        for a in range(len(Chapter_names)):
            wait = WebDriverWait(driver, 10)  # 等待网页加载完成
            wait.until(EC.presence_of_element_located(
                (By.XPATH, "//iframe[@id='iframe']")))
            Video_Fistiframe = driver.find_element(
                By.XPATH, "//iframe[@id='iframe']")
            driver.switch_to.frame(Video_Fistiframe)  # 切换至第一层
            sleep(1)
            Video_names = driver.find_elements(
                By.TAG_NAME, 'iframe')
            if len(Video_names) > 0:
                for iframe in Video_names:
                    Video_data = iframe.get_attribute('data')
                    Video_header.append(iframe.get_attribute('src'))
                    name_pattern = r'"name":"(.*?)"'
                    type_pattern = r'"type":"(.*?)"'
                    name_matches = re.findall(name_pattern, Video_data)
                    type_matches = re.findall(type_pattern, Video_data)
                    for name, type in zip(name_matches, type_matches):
                        if type == '.mp4':
                            Video_name.append(name.encode(
                                'utf-8').decode('unicode_escape'))
            Video_iframe = driver.find_elements(
                By.XPATH, "//iframe[@class='ans-attach-online ans-insertvideo-online']")
            if len(Video_iframe) > 0:
                if os.path.exists(root_path + f"\\学习通章节视频\\{Chapter_names[a]}") == False:
                    os.mkdir(root_path + f"\\学习通章节视频\\{Chapter_names[a]}")
                print(f"{Chapter_names[a]} 发现" +
                      str(len(Video_iframe)) + "个视频")
                for i in range(len(Video_iframe)):
                    driver.switch_to.frame(Video_iframe[i])
                    Video_url.append(driver.find_element(
                        By.XPATH, "//video[@id='video_html5_api']").get_attribute('src'))
                    driver.switch_to.parent_frame()
                if len(Video_name) == len(Video_url):
                    for i in range(len(Video_url)):
                        headers = {
                            'referer': Video_header[i]
                        }
                        response = requests.get(
                            url=Video_url[i], headers=headers, timeout=10, stream=True)
                        if response.status_code != 403:
                            total_size = int(
                                response.headers.get('Content-Length', 0))
                            block_size = 1024
                            file_dir = root_path + \
                                       f"\\学习通章节视频\\{Chapter_names[a]}\\{Video_name[i]}"
                            with open(file_dir, 'wb') as file:
                                with tqdm(total=total_size, unit='B', unit_scale=True) as progress_bar:
                                    for data in response.iter_content(block_size):
                                        file.write(data)
                                        progress_bar.update(len(data))
                        else:
                            print("Get视频地址失败！")
                    print("视频下载完成！")
                    progress_bar.close()
                    Video_name.clear()
                    Video_url.clear()
                    driver.switch_to.default_content()
                    sleep(1)
                    Next_button = driver.find_elements(
                        By.XPATH, "//div[@id='prevNextFocusNext']")
                    if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                        driver.execute_script(
                            "arguments[0].scrollIntoView();", Next_button[0])
                        sleep(1)
                        if len(Next_button) > 0:
                            Next_button[0].click()
                            print("下一节")
                        sleep(10)
                    else:
                        print("爬取视频完成！文件保存路径：" + root_path + "\\学习通章节视频")

                        input("")
                else:
                    print(f"{Chapter_names[a]} 未发现视频！")
            else:
                print(f"{Chapter_names[a]} 未发现视频！")
                Video_name.clear()
                Video_url.clear()
                driver.switch_to.default_content()
                sleep(1)
                Next_button = driver.find_elements(
                    By.XPATH, "//div[@id='prevNextFocusNext']")
                if len(Next_button) > 0 and Next_button[0].get_attribute("style") != "display: none;":
                    driver.execute_script(
                        "arguments[0].scrollIntoView();", Next_button[0])
                    sleep(1)
                    if len(Next_button) > 0:
                        Next_button[0].click()
                        print("下一节")
                    sleep(10)
    else:
        print("获取章节失败！")


def Download_video():  # 爬取当前小节视频
    wait = WebDriverWait(driver, 10)  # 等待网页加载完成
    wait.until(EC.presence_of_element_located(
        (By.XPATH, "//iframe[@id='iframe']")))
    if len(Chapter_names) > 0:
        Video_name = []
        Video_url = []
        Video_header = []
        root_path = os.getcwd()
        if os.path.exists(root_path + r"\学习通章节视频") == False:
            os.mkdir(root_path + r"\学习通章节视频")
        wait = WebDriverWait(driver, 10)  # 等待网页加载完成
        wait.until(EC.presence_of_element_located(
            (By.XPATH, "//iframe[@id='iframe']")))
        Video_Fistiframe = driver.find_element(
            By.XPATH, "//iframe[@id='iframe']")
        driver.switch_to.frame(Video_Fistiframe)  # 切换至第一层
        sleep(1)
        Video_names = driver.find_elements(
            By.TAG_NAME, 'iframe')
        if len(Video_names) > 0:
            for iframe in Video_names:
                Video_data = iframe.get_attribute('data')
                Video_header.append(iframe.get_attribute('src'))
                name_pattern = r'"name":"(.*?)"'
                type_pattern = r'"type":"(.*?)"'
                name_matches = re.findall(name_pattern, Video_data)
                type_matches = re.findall(type_pattern, Video_data)
                for name, type in zip(name_matches, type_matches):
                    if type == '.mp4':
                        Video_name.append(name.encode(
                            'utf-8').decode('unicode_escape'))
        Video_iframe = driver.find_elements(
            By.XPATH, "//iframe[@class='ans-attach-online ans-insertvideo-online']")
        if len(Video_iframe) > 0:
            if os.path.exists(root_path + f"\\学习通章节视频\\{Chapter_names[int(EnterSection_number)]}") == False:
                os.mkdir(
                    root_path + f"\\学习通章节视频\\{Chapter_names[int(EnterSection_number)]}")
            print(f"{Chapter_names[int(EnterSection_number)]} 发现" +
                  str(len(Video_iframe)) + "个视频")
            for i in range(len(Video_iframe)):
                driver.switch_to.frame(Video_iframe[i])
                Video_url.append(driver.find_element(
                    By.XPATH, "//video[@id='video_html5_api']").get_attribute('src'))
                driver.switch_to.parent_frame()
            if len(Video_name) == len(Video_url):
                for i in range(len(Video_url)):
                    headers = {
                        'referer': Video_header[i]
                    }
                    response = requests.get(
                        url=Video_url[i], headers=headers, timeout=10, stream=True)
                    if response.status_code != 403:
                        total_size = int(
                            response.headers.get('Content-Length', 0))
                        block_size = 1024
                        file_dir = root_path + \
                                   f"\\学习通章节视频\\{Chapter_names[int(EnterSection_number)]}\\{Video_name[i]}"
                        with open(file_dir, 'wb') as file:
                            with tqdm(total=total_size, unit='B', unit_scale=True) as progress_bar:
                                for data in response.iter_content(block_size):
                                    file.write(data)
                                    progress_bar.update(len(data))
                    else:
                        print("Get视频地址失败！")
                print("视频下载完成！")
                progress_bar.close()
                Video_name.clear()
                Video_url.clear()
                print("爬取视频完成！文件保存路径：" + root_path + "\\学习通章节视频")
                input("")
            else:
                print(f"{Chapter_names[int(EnterSection_number)]} 未发现视频！")
        else:
            print(f"{Chapter_names[int(EnterSection_number)]} 未发现视频！")
    else:
        print("获取章节失败！")


if __name__ == '__main__':  # Edge驱动更新判断及下载更新文件并解压替换
    Edge_path = get_edge_installation_path()
    if Edge_path != None:
        Edge_version = Get_EdgeVersion()
        print("Edge浏览器已安装，版本号：" + str(Edge_version))
        Edge_Driverpath = Edge_path.replace("\\msedge.exe", "")
        if os.path.isfile(f"{Edge_Driverpath}\\msedgedriver.exe"):
            Driver_version = Get_EdgeDriverVersion(
                f"{Edge_Driverpath}\\msedgedriver.exe")
            print("Edge浏览器驱动已存在，版本号：" + str(Driver_version))
            print("正在对Edge浏览器驱动检查更新中！")
            if Driver_version != Edge_version:
                print("Edge浏览器驱动需要更新，即将下载更新！")
                os.remove(f"{Edge_Driverpath}\\msedgedriver.exe")
                url = 'https://msedgedriver.azureedge.net/' + Edge_version + '/edgedriver_win64.zip'
                response = requests.get(url=url, stream=True)
                file_dir = Edge_Driverpath + "\\edgedriver_win64.zip"
                total_size = int(response.headers.get('Content-Length', 0))
                block_size = 1024
                with open(file_dir, 'wb') as file:
                    with tqdm(total=total_size, unit='B', unit_scale=True) as progress_bar:
                        for data in response.iter_content(block_size):
                            file.write(data)
                            progress_bar.update(len(data))
                print("下载成功！")
                progress_bar.close()
                with zipfile.ZipFile(file_dir, 'r') as zip_ref:
                    zip_ref.extractall(Edge_Driverpath)
                if os.path.isfile(f"{Edge_Driverpath}\\msedgedriver.exe"):
                    os.remove(file_dir)
                    Driver_version = Get_EdgeDriverVersion()
                    if Driver_version == Edge_version:
                        os.remove(file_dir)
                        args = sys.argv[:]
                        subprocess.Popen(args, shell=True)
                        sys.exit()
                    else:
                        print("Edge浏览器驱动更新失败！")
                        input("")
                else:
                    print("Edge浏览器驱动更新失败！")
                    input("")
            else:
                print("Edge浏览器驱动为最新版,无需更新！")
                Username = input("请输入学习通账号：")
                Password = input("请输入学习通密码：")
                webdriver.edge.driver = f"{Edge_Driverpath}\\msedgedriver.exe"
                driver = webdriver.Edge()
                driver.maximize_window()
                if Username != "" and Password != "":
                    Login_state = Login(Username, Password)
                    if Login_state == True:
                        print("登录成功，欢迎使用学习通selenium刷课脚本！")
                        Enterclass = ClassList()
                        if Enterclass == True:
                            print("进入课程成功！")
                            Function()
                        else:
                            print("进入课程失败！")
                            input("")
                    else:
                        print("登录失败！")
                        input("")
        else:
            print("Edge浏览器驱动不存在，即将下载Edge浏览器驱动！")
            url = 'https://msedgedriver.azureedge.net/' + Edge_version + '/edgedriver_win64.zip'
            response = requests.get(url=url, stream=True)
            file_dir = Edge_Driverpath + "\\edgedriver_win64.zip"
            total_size = int(response.headers.get('Content-Length', 0))
            block_size = 1024
            with open(file_dir, 'wb') as file:
                with tqdm(total=total_size, unit='B', unit_scale=True) as progress_bar:
                    for data in response.iter_content(block_size):
                        file.write(data)
                        progress_bar.update(len(data))
            print("下载成功！")
            progress_bar.close()
            with zipfile.ZipFile(file_dir, 'r') as zip_ref:
                zip_ref.extractall(Edge_Driverpath)
            if os.path.isfile(f"{Edge_Driverpath}\\msedgedriver.exe"):
                os.remove(file_dir)
                args = sys.argv[:]
                subprocess.Popen(args, shell=True)
                sys.exit()
            else:
                print("Edge浏览器驱动安装失败！")
                input("")
    else:
        print("Edge浏览器未安装")
